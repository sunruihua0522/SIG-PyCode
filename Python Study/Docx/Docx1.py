from docx import *
from docx.shared import Inches
from docx.enum.table import *
from docx.enum.text import *
file = Document()

#添加标题
file.add_heading('标题',level=2)


#添加段落
p=file.add_paragraph('This is a new paragraph,This is a new paragraph,This is a new paragraph,This is a new paragraph,This is a new paragraph,This is a new paragraph,This is a new paragraph,')
p.add_run('斜体字').italic = True
p.add_run('粗体字').bold = True
p.italic = True
p.font='lisu'
p.line_spacing_rule = WD_LINE_SPACING.DOUBLE


#添加图片
file.add_picture('CombiSmile.png',width = Inches(4.0), height = Inches(2.0))

#添加表格
rowNum = 2
colNum = 5
tb = file.add_table(rows=rowNum, cols=colNum)

for i in range(rowNum):
    for j in range(colNum):
        cell = tb.cell(i,j)
        cell.text='第%d行%d列'%(i+1,j+1)
        cell.vertical_alignment  = WD_TAB_ALIGNMENT.CENTER

file.add_page_break()

#保存文档
file.save('1.docx')